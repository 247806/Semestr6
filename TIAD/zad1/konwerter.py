import tkinter as tk
from tkinter import filedialog

import pandas as pd
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph

def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    if file_path:
        entry_file_path.delete(0, tk.END)
        entry_file_path.insert(0, file_path)

def convert():
    headers, data = read_excel(entry_file_path.get())
    title = entry_title.get()
    add_page = page.get()

    if title == "":
        title = "file"

    column_widths = prepare_col_sizes(headers, data)
    create_docx(headers, data, f"{title}.docx", add_page, column_widths)
    create_pdf(headers, data, f"{title}.pdf", add_page, column_widths)

def read_excel(file_path):
    df = pd.read_excel(file_path, engine="openpyxl").fillna("")  # Zamiana NaN na ""

    df.columns = [f"Kolumna_{i}" if col.startswith("Unnamed") else col for i, col in enumerate(df.columns)]

    headers = df.columns.tolist()
    data = df.values.tolist()
    return headers, data

def prepare_col_sizes(headers, data):
    sizes = []
    for i in range(len(data[0])):
        while len(sizes) <= i:  # Upewniamy się, że istnieje odpowiedni wiersz
            sizes.append([])

        for j, row in enumerate(data):
            while len(sizes[i]) <= j:  # Upewniamy się, że istnieje odpowiednia kolumna
                sizes[i].append(0)
            sizes[i][j] = len(str(data[j][i])) * 0.381  # Średni rozmiar litery w cm
    print(sizes)

    max_col_size = 7.62
    min_col_size = 1.5

    column_widths = []
    for i in range(len(sizes)):
        max_width = max(sizes[i])
        # Ogranicz szerokość kolumny do zakresu [min_col_size, max_col_size]
        column_width = min(max_col_size, max(min_col_size, max_width))
        column_widths.append(column_width)

    return column_widths

def create_docx(headers, data, output_file, add_page, column_widths):
    doc = Document()
    doc.add_heading('Tabela danych', level=1)
    section = doc.sections[0]
    width = section.page_width.mm - section.left_margin.mm - section.right_margin.mm
    print(f"Szerokość strony: {width} punktów")

    page_width_cm = 15.24
    current_width = 0
    start_col = 0  # Indeks pierwszej kolumny w bieżącej tabeli

    for i in range(len(column_widths)):
        # Jeśli dodanie kolejnej kolumny przekroczy dostępną szerokość strony, utwórz nową tabelę
        if current_width + column_widths[i] > page_width_cm:
            # Utwórz tabelę dla kolumn od start_col do i-1
            create_table(doc, headers[start_col:i], [row[start_col:i] for row in data], column_widths[start_col:i])
            start_col = i
            current_width = 0

        current_width += column_widths[i]

    # Utwórz tabelę dla pozostałych kolumn
    if start_col < len(column_widths):
        create_table(doc, headers[start_col:], [row[start_col:] for row in data], column_widths[start_col:])

    if add_page:
        for i, section in enumerate(doc.sections):
            i+=1
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            run = paragraph.add_run()
            run.text = "Strona "
            field_code = f'{i}'  # Pole numerowania stron
            paragraph.add_run(f' {field_code}').bold = True

            # Wyrównanie numeru strony do prawej
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(output_file)
    print(f"Plik DOCX zapisany: {output_file}")

def create_table(doc, headers, data, column_widths):
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=len(headers), style='Table Grid')

    # Ustaw szerokość kolumn
    for i, width in enumerate(column_widths):
        table.columns[i].width = Cm(width)

    # Nagłówki
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    # Dane
    for row in data:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell) if cell else ""

def create_pdf(headers, data, output_file, add_page, column_widths):
    # Tworzenie dokumentu PDF
    pdf = SimpleDocTemplate(output_file, pagesize=letter)
    elements = []
    width, height = letter  # Szerokość i wysokość w punktach
    print(f"Szerokość strony: {width} punktów ({width / 72:.2f} cali)")

    # Dodanie nagłówka
    styles = getSampleStyleSheet()
    elements.append(Paragraph("Tabela danych", styles['Title']))

    page_width_cm = 15.24
    current_width = 0
    start_col = 0  # Indeks pierwszej kolumny w bieżącej tabeli

    for i in range(len(column_widths)):
        # Jeśli dodanie kolejnej kolumny przekroczy dostępną szerokość strony, utwórz nową tabelę
        if current_width + column_widths[i] > page_width_cm:
            # Utwórz tabelę dla kolumn od start_col do i-1
            elements.append(create_table_pdf(headers[start_col:i], [row[start_col:i] for row in data], column_widths[start_col:i]))
            start_col = i
            current_width = 0

        current_width += column_widths[i]

    # Utwórz tabelę dla pozostałych kolumn
    if start_col < len(column_widths):
        elements.append( create_table_pdf(headers[start_col:], [row[start_col:] for row in data], column_widths[start_col:]))

    # Zapisanie dokumentu PDF
    pdf.build(elements, onFirstPage=add_page_number if add_page else None, onLaterPages=add_page_number if add_page else None)
    print(f"Plik PDF zapisany: {output_file}")


def add_page_number(canvas, doc):
    """
    Dodaje numer strony do stopki dokumentu.
    """
    page_num = canvas.getPageNumber()  # Pobierz numer strony
    text = f"Strona {page_num}"

    # Ustaw styl tekstu
    canvas.setFont("Helvetica", 10)

    # Umieść tekst na dole strony, wyśrodkowany
    canvas.drawCentredString(
        x=doc.pagesize[0] / 2,  # Środek strony w poziomie
        y=20,  # 20 punktów od dołu strony
        text=text
    )

def wrap_text(text, max_width):
    """
    Zawija tekst na podstawie maksymalnej szerokości kolumny.
    :param text: Tekst do zawinięcia.
    :param max_width: Maksymalna szerokość kolumny w punktach.
    :param char_width: Przybliżona szerokość jednego znaku w punktach.
    :return: Lista zawiniętych linii tekstu.
    """
    words = text.split()
    lines = []
    current_line = ""

    cm_to_points = lambda cm: cm * 28.35
    char_width = cm_to_points(0.381)

    for word in words:
        # Sprawdź, czy dodanie kolejnego słowa przekroczy maksymalną szerokość
        if len(current_line) * char_width + len(word) * char_width <= max_width:
            current_line += word + " "
        else:
            lines.append(current_line.strip())
            current_line = word + " "

    # Dodaj ostatnią linię
    if current_line:
        lines.append(current_line.strip())

    return lines

def create_table_pdf(headers, data, column_widths):
    # Przygotowanie danych do tabeli (nagłówki + dane)
    table_data = [headers]

    cm_to_points = lambda cm: cm * 28.35
    for i, size in enumerate(column_widths):
        column_widths[i] = cm_to_points(size)

    for row in data:
        wrapped_row = []
        for i, cell in enumerate(row):
            # Zawijanie tekstu dla każdej komórki
            wrapped_text = wrap_text(str(cell), column_widths[i])
            wrapped_row.append("\n".join(wrapped_text))  # Łączymy linie za pomocą \n
        table_data.append(wrapped_row)

    # Tworzenie tabeli
    table = Table(table_data, colWidths=column_widths)

    # Styl tabeli
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),  # Kolor tła dla nagłówków
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),  # Kolor tekstu dla nagłówków
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),# Wyśrodkowanie tekstu
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Czcionka dla nagłówków
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),  # Dodatkowy padding dla nagłówków
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),  # Kolor tła dla danych
        ('GRID', (0, 0), (-1, -1), 1, colors.black),  # Obramowanie tabeli
    ]))

    # Dodanie tabeli do dokumentu
    return table


# Tworzenie GUI
root = tk.Tk()
root.title("Konwerter XLSX do DOCX/PDF")

frame = tk.Frame(root, padx=10, pady=10)
frame.pack(padx=10, pady=10)

tk.Label(frame, text="Wybierz plik .xlsx:").grid(row=0, column=0, sticky="w")
entry_file_path = tk.Entry(frame, width=50)
entry_file_path.grid(row=0, column=1)
tk.Button(frame, text="Wybierz", command=select_file).grid(row=0, column=2, padx=10)

tk.Label(frame, text="Tytuł dokumentu:").grid(row=1, column=0, sticky="w", pady=10)
entry_title = tk.Entry(frame, width=50)
entry_title.grid(row=1, column=1, columnspan=2, sticky="w")

page = tk.BooleanVar()
check_page = tk.Checkbutton(frame, text="Numeruj strony", variable=page)
check_page.grid(row=2, column=0, pady=5)

# Przycisk konwersji
tk.Button(frame, text="Konwertuj", command=convert).grid(row=3, column=1, pady=10)

root.mainloop()


