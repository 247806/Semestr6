from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def prepare_data(headers, data, widths):
    page_width_cm = 16.51
    current_width = 0
    start_col = 0
    data1 = []
    data2 = []
    data3 = []

    for i in range(len(widths)):
        if current_width + widths[i] > page_width_cm:
            data1.append(headers[start_col:i])
            data2.append([row[start_col:i] for row in data])
            data3.append(widths[start_col:i])
            start_col = i
            current_width = 0

        current_width += widths[i]

    if start_col < len(widths):
        data1.append(headers[start_col:])
        data2.append([row[start_col:] for row in data])
        data3.append(widths[start_col:])

    return data1, data2, data3

def create_docx(headers, data, output_file, add_page, column_widths, align, heading, table, col_delete):
    doc = Document()
    doc.add_heading(heading, level=1)

    for i in range(len(col_delete) - 1, -1, -1):
        if not col_delete[i].get():
            del headers[i]
            for row in data:
                del row[i]
            del column_widths[i]

    headers_line, data_line, widths = prepare_data(headers, data, column_widths)
    if table:
        for i in range(len(headers_line)):
            create_table(doc, headers_line[i], data_line[i], widths[i], align)
    else:
        create_paragraphs(doc, headers, data, align)


    if add_page:
        for i, section in enumerate(doc.sections):
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = paragraph.add_run("Strona ")

            fldSimple = OxmlElement('w:fldSimple')
            fldSimple.set(qn('w:instr'), "PAGE")
            fldRun = OxmlElement('w:r')
            fldSimple.append(fldRun)
            run._r.append(fldSimple)

    doc.save(output_file)
    # print(f"Plik DOCX zapisany: {output_file}")

def create_table(doc, headers, data, column_widths, align):
    align_map = {
        "Do lewej": WD_ALIGN_PARAGRAPH.LEFT,
        "Do środka": WD_ALIGN_PARAGRAPH.CENTER,
        "Do prawej": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    doc.add_paragraph()
    table = doc.add_table(rows=0, cols=len(headers), style='Table Grid')

    for i, width in enumerate(column_widths):
        table.columns[i].width = Cm(width)

    # Nagłówki
    hdr_cells = table.add_row().cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        paragraph = hdr_cells[i].paragraphs[0]
        paragraph.alignment = align_map[align]
        run = paragraph.runs[0]
        run.bold = True

        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '2E75B6')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

    # Dane
    for row in data:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell) if cell else ""
            row_cells[i].paragraphs[0].alignment = align_map[align]

def create_paragraphs(doc, headers, data, align):
    align_map = {
        "Do lewej": WD_ALIGN_PARAGRAPH.LEFT,
        "Do środka": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "Do prawej": WD_ALIGN_PARAGRAPH.RIGHT,
    }

    doc.add_paragraph()

    for row_idx, row in enumerate(data):
        lines = [f"{headers[col_idx]}: {row[col_idx]}" for col_idx in range(len(headers))]
        paragraph_text = f"{row_idx + 1}.\n" + '\n'.join(lines)

        paragraph = doc.add_paragraph(paragraph_text)
        if align in align_map:
            paragraph.alignment = align_map[align]

        doc.add_paragraph()



